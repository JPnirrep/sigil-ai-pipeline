"""
KLEIA-UP Book — KDP DOCX Template Parser
Parse Amazon KDP-style DOCX templates into semantic XHTML + CSS
"""

import os, re
from dataclasses import dataclass, field
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt


# ── KDP Template Metadata ──

@dataclass
class KDPMetadata:
    """Detected KDP template parameters"""
    trim_width: float = 6.0        # inches
    trim_height: float = 9.0
    margin_top: float = 0.75
    margin_bottom: float = 0.85
    margin_inner: float = 0.9      # gutter (binding side)
    margin_outer: float = 0.6
    bleed: float = 0.125
    detected: bool = False


# ── KDP Style Map ──

KDP_STYLE_MAP = {
    "Title":                     {"tag": "h1", "class": "book-title"},
    "Subtitle":                  {"tag": "h2", "class": "book-subtitle"},
    "Author":                    {"tag": "p",  "class": "author"},
    "Heading 1":                 {"tag": "h1", "class": "chapter-title"},
    "Heading 2":                 {"tag": "h2", "class": "section-title"},
    "Heading 3":                 {"tag": "h3", "class": "sub-section"},
    "Body Text":                 {"tag": "p",  "class": "body-text"},
    "Body Text First Paragraph": {"tag": "p",  "class": "first-para"},
    "Normal":                    {"tag": "p",  "class": "body-text"},
    "Copyright Page":            {"tag": "section", "class": "copyright"},
    "Image":                     {"tag": "figure", "class": "image"},
    "TOC Heading":               {"tag": "h1", "class": "toc-title"},
    "Dedication":                {"tag": "section", "class": "dedication"},
    "Epigraph":                  {"tag": "blockquote", "class": "epigraph"},
    "Block Quote":               {"tag": "blockquote"},
}

# Reverse: XHTML class → print CSS behaviors
CLASS_PRINT_RULES = {
    "book-title":    {"page-break": "after", "text-align": "center"},
    "chapter-title": {"page-break": "before", "page": "chapter", "string-set": "chapter-title content()"},
    "first-para":    {"text-indent": "0"},
    "body-text":     {"text-indent": "1.5em", "hyphens": "auto"},
    "copyright":     {"page-break": "before", "font-size": "9pt"},
    "dedication":    {"page-break": "before"},
}


# ── Parser ──

@dataclass
class ParsedChapter:
    title: str
    elements: list = field(default_factory=list)  # list of (tag, class_, content)

@dataclass
class ParsedBook:
    title: str = ""
    subtitle: str = ""
    author: str = ""
    metadata: KDPMetadata = field(default_factory=KDPMetadata)
    chapters: list = field(default_factory=list)
    front_matter: list = field(default_factory=list)
    raw_styles_detected: list = field(default_factory=list)


def parse_docx(path: str) -> ParsedBook:
    """Parse a KDP-format DOCX into a structured ParsedBook"""
    doc = Document(path)
    book = ParsedBook()
    book.raw_styles_detected = _detect_styles(doc)
    book.metadata = _detect_template(doc)

    current_chapter = None

    for p in doc.paragraphs:
        style_name = p.style.name
        text = p.text.strip()
        if not text:
            continue

        mapping = KDP_STYLE_MAP.get(style_name)
        if not mapping:
            mapping = _fuzzy_match(style_name)

        tag = mapping["tag"] if mapping else "p"
        cls = mapping.get("class", "")

        # Route into book structure
        if cls == "book-title":
            book.title = text
        elif cls == "book-subtitle":
            book.subtitle = text
        elif cls == "author":
            book.author = text
        elif cls == "chapter-title":
            if current_chapter:
                book.chapters.append(current_chapter)
            current_chapter = ParsedChapter(title=text)
        else:
            if current_chapter:
                current_chapter.elements.append((tag, cls, text))
            else:
                book.front_matter.append((tag, cls, text))

    if current_chapter:
        book.chapters.append(current_chapter)

    return book


def _detect_styles(doc: Document) -> list:
    """List all paragraph styles used in the document"""
    seen = set()
    for p in doc.paragraphs:
        if p.text.strip():
            seen.add(p.style.name)
    return sorted(seen)


def _detect_template(doc: Document) -> KDPMetadata:
    """Detect KDP template parameters from page setup"""
    meta = KDPMetadata()
    try:
        s = doc.sections[0]
        meta.trim_width = round(s.page_width / 914400 * 2.54 / 2.54, 1)      # EMU → inches
        meta.trim_height = round(s.page_height / 914400 * 2.54 / 2.54, 1)
        meta.margin_top = round(s.top_margin / 914400 * 2.54 / 2.54, 3)
        meta.margin_bottom = round(s.bottom_margin / 914400 * 2.54 / 2.54, 3)
        meta.margin_inner = round(s.left_margin / 914400 * 2.54 / 2.54, 3)
        meta.margin_outer = round(s.right_margin / 914400 * 2.54 / 2.54, 3)
        meta.detected = True
    except Exception:
        pass
    return meta


def _fuzzy_match(style_name: str) -> dict:
    """Fallback: map unknown style names by pattern"""
    name = style_name.lower()
    if "heading" in name or "h1" in name or "chapter" in name:
        return {"tag": "h1", "class": "chapter-title"}
    if "body" in name or "text" in name or "normal" in name:
        return {"tag": "p", "class": "body-text"}
    if "title" in name:
        return {"tag": "h1", "class": "book-title"}
    if "quote" in name or "citation" in name:
        return {"tag": "blockquote"}
    return {"tag": "p", "class": ""}


def _classify_front(text: str) -> str:
    """Classify front matter paragraph into section type."""
    t = text.strip().lower()
    if not t: return "generic"
    if t.startswith("table des mati"): return "toc-title"
    if t.startswith("introduction"): return "intro-title"
    if any(w in t for w in ["copyright", "tous droits", "isbn", "©"]): return "copyright"
    if t.startswith("dédicace") or t.startswith("a mes"): return "dedication"
def _highlight_numbers(text: str) -> str:
    """Smart number highlighting: key numbers only, not addresses/dates."""
    import re
    escaped = _escape(text)

    # 1) Leading number at paragraph start: "1. La manifestation"
    escaped = re.sub(r'^(\d+)(\.)', r'<span class="num">\1</span>\2', escaped)

    # 2) Leading number at line start after tag
    escaped = re.sub(r'^(\d+)(\s)', r'<span class="num">\1</span>\2', escaped)

    # 3) Number + period + space mid-line (enumerations): " 2. Le nombre"
    escaped = re.sub(r'(\s)(\d{1,2})(\.)(\s)', r'\1<span class="num">\2</span>\3\4', escaped)

    # 4) Small inline numbers (1-2 digits) between spaces or parens
    # Excludes 3+ digit numbers (addresses) and year patterns (19xx, 20xx)
    escaped = re.sub(
        r'(?<=[\s(])(\d{1,2})(?=[\s,).;])',
        r'<span class="num">\1</span>', escaped
    )

    return escaped
def render_xhtml(book: ParsedBook) -> str:
    """Render ParsedBook as smart XHTML with structured front matter and number highlighting."""
    parts = ["<body>"]

    # ── Title page ──
    parts.append('<section class="title-page">')
    if book.title:
        parts.append(f'<h1 class="book-title">{_escape(book.title)}</h1>')
    if book.subtitle:
        parts.append(f'<h2 class="book-subtitle">{_escape(book.subtitle)}</h2>')
    if book.author:
        parts.append(f'<p class="author">{_escape(book.author)}</p>')
    parts.append('</section>')

    # ── Front matter sections ──
    current_section = None
    sections_created = set()

    for tag, cls, text in book.front_matter:
        ft = _classify_front(text)

        # TOC title — open TOC section (once)
        if ft == "toc-title" and "toc" not in sections_created:
            if current_section: parts.append('</section>')
            parts.append('<section class="front-matter-toc">')
            parts.append('<h2 class="toc-title">Table des Matières</h2>')
            current_section = "toc"
            sections_created.add("toc")
            continue

        # If we're in TOC: check if this is still a TOC entry
        if current_section == "toc":
            is_toc_entry = bool(re.match(r'^[IVXLCDM]+\.\s', text)) or \
                           bool(re.match(r'^Introduction\s+\d+', text)) or \
                           bool(re.match(r'^\d+\.\s', text))
            if is_toc_entry:
                attr = f' class="{cls}"' if cls else ""
                parts.append(f'<{tag}{attr}>{_highlight_numbers(text)}</{tag}>')
                continue
            # Not a TOC entry: close TOC, then fall through to render this paragraph
            parts.append('</section>')
            current_section = None
            # Fall through to generic rendering below

        # Introduction / copyright (only when NOT in TOC)
        if ft == "intro-title" and "intro" not in sections_created:
            if current_section: parts.append('</section>')
            parts.append('<section class="front-matter-intro">')
            parts.append('<h2 class="intro-title">Introduction</h2>')
            current_section = "intro"
            sections_created.add("intro")
            continue

        if ft == "copyright" and "copyright" not in sections_created:
            if current_section: parts.append('</section>')
            parts.append('<section class="front-matter-copyright">')
            current_section = "copyright"
            sections_created.add("copyright")

        # Generic: render with current section or standalone
        attr = f' class="{cls}"' if cls else ""
        parts.append(f'<{tag}{attr}>{_highlight_numbers(text)}</{tag}>')

    # ── Chapters ──
    for ch in book.chapters:
        parts.append(f'<section class="chapter" epub:type="chapter">')
        parts.append(f'<h1 class="chapter-title">{_escape(ch.title)}</h1>')
        for tag, cls, text in ch.elements:
            attr = f' class="{cls}"' if cls else ""
            parts.append(f'<{tag}{attr}>{_highlight_numbers(text)}</{tag}>')
        parts.append('</section>')

    parts.append("</body>")
    return "\n".join(parts)


def render_css_epub(book: ParsedBook) -> str:
    """Generate EPUB reflowable CSS from book metadata"""
    template = book.metadata
    return f"""/* KLEIA-UP Book — EPUB Reflowable Theme */
body {{
  font-family: Georgia, "Times New Roman", serif;
  line-height: 1.5;
  margin: 0;
  padding: 0;
}}

h1.book-title {{
  text-align: center;
  font-size: 2em;
  margin-top: 20%;
  margin-bottom: 0.5em;
}}

h2.book-subtitle {{
  text-align: center;
  font-size: 1.4em;
  font-style: italic;
  font-weight: normal;
}}

p.author {{
  text-align: center;
  font-size: 1.2em;
  margin-top: 2em;
}}

h1.chapter-title {{
  text-align: left;
  font-size: 1.6em;
  margin-top: 2em;
  margin-bottom: 1em;
  page-break-before: always;
}}

section.chapter p.body-text {{
  text-indent: 1.5em;
  margin: 0 0 0.5em 0;
  hyphens: auto;
}}

section.chapter p.first-para {{
  text-indent: 0;
  margin: 0 0 0.5em 0;
}}


blockquote {{
  margin: 1em 2em;
  font-style: italic;
}}

section.copyright {{
  font-size: 0.8em;
  text-align: center;
  margin-top: 10%;
}}

/* ── Front matter sections ── */
section.title-page {{
  page-break-after: always;
  text-align: center;
  padding-top: 20%;
}}

section.front-matter-toc {{
  page-break-before: always;
}}

section.front-matter-intro {{
  page-break-before: always;
}}

section.front-matter-copyright {{
  page-break-before: always;
  font-size: 0.8em;
  text-align: center;
}}

/* ── Number highlighting ── */
span.num {{
  font-weight: bold;
  font-size: 1.1em;
  color: #8b4513;
}}
"""




def render_css_print(book: ParsedBook) -> str:
    """Generate print CSS with Paged Media from book metadata"""
    m = book.metadata
    return f"""/* KLEIA-UP Book — Print PDF Theme */
@page {{
  size: {m.trim_width}in {m.trim_height}in;
  margin: {m.margin_top}in {m.margin_outer}in {m.margin_bottom}in {m.margin_inner}in;
  bleed: {m.bleed}in;
}}

@page :left {{
  margin-left: {m.margin_inner}in;
  margin-right: {m.margin_outer}in;
  @top-left {{
    content: string(book-title);
    font-size: 8pt;
    font-style: italic;
  }}
  @bottom-left {{
    content: counter(page);
    font-size: 8pt;
  }}
}}

@page :right {{
  margin-left: {m.margin_outer}in;
  margin-right: {m.margin_inner}in;
  @top-right {{
    content: string(chapter-title);
    font-size: 8pt;
    font-style: italic;
  }}
  @bottom-right {{
    content: counter(page);
    font-size: 8pt;
  }}
}}

@page chapter:first {{
  @top-left {{ content: none; }}
  @top-right {{ content: none; }}
}}

body {{
  font-family: "Times New Roman", Georgia, serif;
  font-size: 11pt;
  line-height: 1.3;
  color: #000;
}}

h1.book-title {{
  string-set: book-title content(text);
  text-align: center;
  font-size: 24pt;
  margin-top: 30%;
  page-break-after: always;
}}

h1.chapter-title {{
  string-set: chapter-title content(text);
  page: chapter;
  page-break-before: right;
  font-size: 16pt;
  text-align: center;
  margin-top: 20%;
  margin-bottom: 1.5em;
}}

p.body-text {{
  text-indent: 1.5em;
  margin: 0 0 0.3em 0;
  hyphens: auto;
  widows: 2;
  orphans: 2;
}}

p.first-para {{
  text-indent: 0;
  margin: 0 0 0.3em 0;
}}

/* ── Front matter sections ── */
section.title-page {{
  page-break-after: always;
  text-align: center;
  padding-top: 30%;
}}

section.front-matter-toc {{
  page-break-before: left;
}}

section.front-matter-intro {{
  page-break-before: left;
}}

/* ── Number highlighting ── */
span.num {{
  font-weight: bold;
  font-size: 1.1em;
  color: #8b4513;
}}
"""


def _escape(text: str) -> str:
    """Basic XML escape"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
